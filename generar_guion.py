"""Genera el documento Word con el guion de exposición del proyecto."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x36, 0x5C)
    return h


def add_paragraph(doc, text: str, bold: bool = False, italic: bool = False, size: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            label, body = item
            run = p.add_run(label)
            run.bold = True
            p.add_run(": " + body)
        else:
            p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, header_color: str = "1F365C"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], header_color)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            cells[c_idx].paragraphs[0].add_run(str(val))
    return table


def add_qa(doc, qa_list):
    for q, a in qa_list:
        p = doc.add_paragraph()
        run = p.add_run("P. ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x36, 0x5C)
        run2 = p.add_run(q)
        run2.bold = True
        a_p = doc.add_paragraph()
        ar = a_p.add_run("R. ")
        ar.bold = True
        ar.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        a_p.add_run(a)


doc = Document()

# Estilos por defecto
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# =====================================================================
# PORTADA
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Guion de exposición")
tr.bold = True
tr.font.size = Pt(28)
tr.font.color.rgb = RGBColor(0x1F, 0x36, 0x5C)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Simulador de Sistemas Operativos\nAlgoritmos de planificación de CPU y reemplazo de páginas")
sr.font.size = Pt(16)
sr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Proyecto final | Sistemas Operativos | UDEM | Primavera 2026")
mr.italic = True
mr.font.size = Pt(12)

doc.add_paragraph()
intro = doc.add_paragraph()
ir = intro.add_run(
    "Este documento es el guion completo para defender el demo del simulador. "
    "Incluye la teoría detrás de cada algoritmo, ejemplos numéricos resueltos, "
    "comparativas, descripción técnica del demo y un banco amplio de preguntas y "
    "respuestas que podría hacer el profesor o un compañero. La idea es que, leyendo "
    "este archivo, se pueda responder con seguridad a cualquier pregunta sobre el "
    "comportamiento de los algoritmos y la implementación del proyecto."
)
ir.font.size = Pt(11)

doc.add_page_break()

# =====================================================================
# 0. ÍNDICE
# =====================================================================
add_heading(doc, "Índice", level=1)
add_bullets(doc, [
    "1. Presentación del demo",
    "2. Conceptos previos de procesos",
    "3. Métricas de planificación",
    "4. Algoritmos de planificación de CPU",
    "5. Conceptos previos de memoria virtual",
    "6. Algoritmos de reemplazo de páginas",
    "7. Comparativas rápidas",
    "8. Arquitectura técnica del demo",
    "9. Banco de preguntas y respuestas",
    "10. Errores y trampas comunes",
    "11. Cierre y conclusiones",
])
doc.add_page_break()

# =====================================================================
# 1. PRESENTACIÓN DEL DEMO
# =====================================================================
add_heading(doc, "1. Presentación del demo", level=1)
add_paragraph(doc,
    "El proyecto es una aplicación web interactiva que permite visualizar, paso a paso, "
    "los conceptos centrales de un sistema operativo: gestión de procesos e hilos, "
    "planificación de CPU, paginación de memoria y algoritmos de reemplazo de páginas. "
    "Está construida con un enfoque mobile-first y combina una guía interactiva por "
    "capítulos con un modo libre que permite configurar simulaciones completas."
)
add_heading(doc, "1.1 Tres zonas del producto", level=2)
add_bullets(doc, [
    ("Bienvenida", "portada del simulador con accesos directos a la guía o al modo libre."),
    ("Guía interactiva", "tutorial paso a paso dividido en seis capítulos con analogías, animaciones y mini simulaciones."),
    ("Modo libre", "seis módulos independientes (procesos, planificación, memoria, reemplazo, métricas y comparación) donde se pueden configurar y ejecutar simulaciones completas."),
])

add_heading(doc, "1.2 Tecnologías principales", level=2)
add_table(doc,
    ["Categoría", "Herramienta"],
    [
        ["Lenguaje", "TypeScript"],
        ["UI", "React 19"],
        ["Empaquetador", "Vite"],
        ["Estilos", "Tailwind CSS"],
        ["Enrutamiento", "React Router DOM"],
        ["Estado global", "Zustand con persist"],
        ["Animaciones", "Framer Motion"],
        ["Gráficas", "Recharts"],
        ["Hospedaje", "Firebase Hosting"],
        ["CI/CD", "GitHub Actions"],
    ])

add_heading(doc, "1.3 Qué se va a demostrar", level=2)
add_numbered(doc, [
    "Crear procesos manualmente o por lote y visualizar su estado.",
    "Elegir un algoritmo de planificación de CPU, configurarlo y ejecutar la simulación con un diagrama de Gantt en vivo.",
    "Configurar la memoria física, el tamaño de página y el número de marcos.",
    "Lanzar una simulación de reemplazo de páginas con la cadena de referencias generada a partir de los procesos.",
    "Ver el panel de métricas que consolida todo: tiempos promedio, utilización de CPU, tasa de fallos.",
    "Comparar varios algoritmos sobre la misma carga en gráficas lado a lado.",
])
doc.add_page_break()

# =====================================================================
# 2. CONCEPTOS PREVIOS DE PROCESOS
# =====================================================================
add_heading(doc, "2. Conceptos previos de procesos", level=1)
add_paragraph(doc,
    "Un proceso es un programa en ejecución. El sistema operativo lo describe con un "
    "PCB (Process Control Block) que contiene la información necesaria para administrarlo."
)
add_heading(doc, "2.1 Atributos importantes del PCB", level=2)
add_bullets(doc, [
    ("PID", "identificador único del proceso."),
    ("Estado", "New, Ready, Running, Waiting o Terminated."),
    ("Tiempo de llegada (arrival time)", "instante en el que el proceso entra al sistema."),
    ("Ráfaga de CPU (burst time)", "tiempo total que necesita ejecutarse en CPU."),
    ("Prioridad", "valor numérico que indica preferencia. En este simulador, número menor significa prioridad mayor."),
    ("Páginas", "número de páginas virtuales que usa el proceso, requerido para los módulos de memoria y reemplazo."),
])

add_heading(doc, "2.2 Estados de un proceso", level=2)
add_table(doc,
    ["Estado", "Significado"],
    [
        ["New", "Recién creado, todavía no admitido al sistema."],
        ["Ready", "Listo en cola, esperando un núcleo libre."],
        ["Running", "Ejecutándose actualmente en la CPU."],
        ["Waiting", "Bloqueado esperando un evento, típicamente E/S."],
        ["Terminated", "Finalizado, libera recursos."],
    ])

add_heading(doc, "2.3 Tipos de planificadores", level=2)
add_table(doc,
    ["Nivel", "Decide cuándo"],
    [
        ["Largo plazo", "Admitir o no procesos al sistema."],
        ["Mediano plazo", "Suspender o reanudar procesos (swap)."],
        ["Corto plazo", "Asignar la CPU al siguiente proceso de la cola de listos."],
    ])
add_paragraph(doc,
    "Este proyecto se concentra en el planificador de corto plazo. Las dos categorías "
    "más importantes a recordar son:"
)
add_bullets(doc, [
    ("No expropiativo (non-preemptive)", "una vez que un proceso recibe la CPU la mantiene hasta terminar o bloquearse."),
    ("Expropiativo (preemptive)", "el sistema puede quitarle la CPU al proceso en ejecución cuando ocurre un evento como la llegada de un proceso de mayor prioridad o el fin del quantum."),
])

add_heading(doc, "2.4 Hilos", level=2)
add_paragraph(doc,
    "Un hilo es la unidad mínima de ejecución dentro de un proceso. Comparten memoria "
    "y recursos del proceso pero tienen su propio contexto de ejecución (PC, registros, "
    "pila). El simulador permite crear hilos por proceso y simular un fork (duplicación). "
    "Los algoritmos de planificación trabajan a nivel de hilo cuando un proceso tiene varios."
)
doc.add_page_break()

# =====================================================================
# 3. MÉTRICAS
# =====================================================================
add_heading(doc, "3. Métricas de planificación", level=1)
add_paragraph(doc, "Para un proceso i se definen los siguientes tiempos:")
add_bullets(doc, [
    ("Tiempo de llegada (Ai)", "cuándo entra al sistema."),
    ("Tiempo de ráfaga (Bi)", "CPU que requiere."),
    ("Tiempo de inicio (Si)", "primera vez que la CPU lo atiende."),
    ("Tiempo de finalización (Fi)", "instante en que termina."),
])

add_heading(doc, "3.1 Métricas derivadas", level=2)
add_table(doc,
    ["Métrica", "Fórmula", "Significado"],
    [
        ["Tiempo de retorno (Turnaround)", "Ti = Fi - Ai", "Cuánto tardó desde que llegó hasta que terminó."],
        ["Tiempo de espera (Waiting)", "Wi = Ti - Bi", "Tiempo total en cola de listos."],
        ["Tiempo de respuesta (Response)", "Ri = Si - Ai", "Cuánto tardó en ser atendido por primera vez."],
    ])
add_paragraph(doc,
    "Los promedios se obtienen sumando los valores individuales y dividiendo entre n."
)

add_heading(doc, "3.2 Otras métricas usadas en el simulador", level=2)
add_bullets(doc, [
    ("Throughput", "procesos terminados por unidad de tiempo."),
    ("Utilización de CPU", "fracción del tiempo en que la CPU está ocupada (idealmente cercana al 100 %)."),
    ("Cambios de contexto", "cuántas veces se reasignó la CPU a un proceso distinto. Cada cambio cuesta tiempo real, por eso menos cambios suele ser mejor."),
    ("Tasa de fallos de página", "F/N, donde F son los fallos y N las referencias totales."),
    ("Tasa de aciertos", "1 - F/N."),
])
doc.add_page_break()

# =====================================================================
# 4. ALGORITMOS DE PLANIFICACIÓN DE CPU
# =====================================================================
add_heading(doc, "4. Algoritmos de planificación de CPU", level=1)
add_paragraph(doc,
    "El simulador implementa ocho algoritmos en src/engine/scheduling/. Cada uno recibe la "
    "lista de procesos y la configuración (quantum, niveles, núcleos) y produce una línea de "
    "tiempo con métricas por proceso, promedios, utilización y cambios de contexto."
)

# --- 4.1 FCFS ---
add_heading(doc, "4.1 FCFS - First Come First Served", level=2)
add_bullets(doc, [
    ("Tipo", "no expropiativo."),
    ("Idea", "el primero que llega es el primero en ser atendido. Cola FIFO simple."),
])
add_paragraph(doc, "Ventajas:", bold=True)
add_bullets(doc, [
    "Muy simple de implementar.",
    "Justo en el sentido cronológico: nadie 'se cuela'.",
])
add_paragraph(doc, "Desventajas:", bold=True)
add_bullets(doc, [
    "Efecto convoy: un proceso largo retrasa a todos los cortos que llegaron después, lo que dispara el tiempo de espera promedio.",
    "Mal tiempo de respuesta para tareas interactivas.",
])
add_paragraph(doc, "Ejemplo numérico:", bold=True)
add_table(doc,
    ["Proceso", "Llegada", "Ráfaga"],
    [["P1", "0", "7"], ["P2", "2", "4"], ["P3", "4", "1"]])
add_paragraph(doc,
    "Orden de ejecución: P1 (0-7), P2 (7-11), P3 (11-12). "
    "Esperas: P1 = 0, P2 = 5, P3 = 7. Promedio = 4. "
    "Nótese que P3, que solo necesitaba 1 unidad, tuvo que esperar 7. Eso es el efecto convoy."
)

# --- 4.2 SJF ---
add_heading(doc, "4.2 SJF - Shortest Job First", level=2)
add_bullets(doc, [
    ("Tipo", "no expropiativo en este simulador."),
    ("Idea", "cuando la CPU se libera, escoge entre los procesos listos el que tenga la menor ráfaga total."),
])
add_paragraph(doc, "Ventajas:", bold=True)
add_bullets(doc, [
    "Es óptimo en tiempo de espera promedio cuando todos los procesos están disponibles al mismo tiempo. Esto se demuestra con un argumento de intercambio: si un proceso largo va antes que uno corto, intercambiarlos reduce la suma de esperas.",
])
add_paragraph(doc, "Desventajas:", bold=True)
add_bullets(doc, [
    "Requiere conocer la ráfaga de antemano. En sistemas reales se estima con un promedio exponencial.",
    "Riesgo de inanición (starvation): un proceso largo puede no ejecutarse nunca si siempre llegan cortos.",
    "Mitigación: aging (envejecimiento), aumentar la prioridad con el tiempo de espera.",
])

# --- 4.3 SRTF ---
add_heading(doc, "4.3 SRTF - Shortest Remaining Time First", level=2)
add_bullets(doc, [
    ("Tipo", "expropiativo. Es la versión preemptiva de SJF."),
    ("Idea", "en cada instante ejecuta el proceso cuyo tiempo restante sea menor. Si llega uno con menos ráfaga restante que el actual, lo expropia."),
])
add_paragraph(doc, "Ventajas:", bold=True)
add_bullets(doc, [
    "Tiempo de espera promedio aún mejor que SJF en cargas con llegadas dinámicas.",
])
add_paragraph(doc, "Desventajas:", bold=True)
add_bullets(doc, [
    "Más cambios de contexto.",
    "Misma posibilidad de inanición que SJF.",
    "Necesita conocer o estimar el tiempo restante.",
])

# --- 4.4 HRRN ---
add_heading(doc, "4.4 HRRN - Highest Response Ratio Next", level=2)
add_bullets(doc, [
    ("Tipo", "no expropiativo."),
    ("Idea", "cuando la CPU se libera, calcula para cada proceso listo el response ratio:"),
])
add_paragraph(doc,
    "    RR = (tiempo de espera + ráfaga) / ráfaga",
    italic=True
)
add_paragraph(doc,
    "y escoge el de mayor RR. Es decir, mezcla 'qué tan corto es el proceso' con 'qué tanto lleva esperando'."
)
add_paragraph(doc, "Ventajas:", bold=True)
add_bullets(doc, [
    "Combina lo mejor de SJF (favorece cortos) y FCFS (favorece a quien lleva mucho esperando).",
    "Evita la inanición: a mayor espera mayor RR, así que tarde o temprano cualquier proceso se ejecuta.",
])
add_paragraph(doc, "Desventajas:", bold=True)
add_bullets(doc, [
    "Requiere recalcular el ratio en cada decisión.",
    "Sigue necesitando conocer la ráfaga.",
])

# --- 4.5 Round Robin ---
add_heading(doc, "4.5 Round Robin (RR)", level=2)
add_bullets(doc, [
    ("Tipo", "expropiativo."),
    ("Idea", "cada proceso recibe la CPU durante un quantum q. Si no termina, se manda al final de la cola y entra el siguiente."),
])
add_paragraph(doc, "Ventajas:", bold=True)
add_bullets(doc, [
    "Excelente tiempo de respuesta, ideal para sistemas interactivos.",
    "No hay inanición: todos los procesos avanzan en el peor caso cada (n-1) * q unidades.",
])
add_paragraph(doc, "Desventajas:", bold=True)
add_bullets(doc, [
    "Muy sensible al valor del quantum.",
    "Quantum muy chico: muchos cambios de contexto, sobrecarga.",
    "Quantum muy grande: degenera en FCFS.",
    "Tiempo de retorno promedio puede ser peor que en SJF/SRTF.",
])
add_paragraph(doc,
    "Regla práctica: el quantum se elige un poco mayor que el tiempo típico de una "
    "interacción de usuario. En sistemas reales suele estar entre 10 y 100 ms."
)

# --- 4.6 Prioridad expropiativa ---
add_heading(doc, "4.6 Prioridad expropiativa", level=2)
add_bullets(doc, [
    ("Tipo", "expropiativo."),
    ("Idea", "cada proceso tiene una prioridad. La CPU siempre ejecuta el proceso listo con mayor prioridad. Si llega uno con prioridad mayor que el actual, lo expropia."),
])
add_paragraph(doc,
    "Convención del simulador: número más bajo significa prioridad más alta (1 es más prioritario que 5)."
)
add_paragraph(doc, "Ventajas:", bold=True)
add_bullets(doc, [
    "Permite reflejar la importancia real de cada tarea: tiempo real, sistema, usuario, batch.",
])
add_paragraph(doc, "Desventajas:", bold=True)
add_bullets(doc, [
    "Inanición de procesos de baja prioridad. Se mitiga con aging.",
    "Inversión de prioridad: un proceso de alta prioridad puede quedar bloqueado esperando un recurso que tiene uno de baja prioridad. Se mitiga con priority inheritance.",
])

# --- 4.7 Cola multinivel ---
add_heading(doc, "4.7 Cola multinivel", level=2)
add_bullets(doc, [
    ("Tipo", "depende del nivel; cada cola usa su propio algoritmo."),
    ("Idea", "dividir los procesos en clases fijas (por ejemplo: sistema, interactivos, batch), cada una con su propia cola y su propio algoritmo (RR para interactivos, FCFS para batch, etc.)."),
])
add_paragraph(doc,
    "En este simulador hay tres colas y la prioridad del proceso decide a cuál entra: "
    "prioridad 1 va a la cola 0 (más prioritaria), prioridad 2 a la cola 1, y "
    "el resto a la cola 2 (menos prioritaria). Entre colas hay prioridad estricta: "
    "la cola superior se atiende primero. Las colas 0 y 1 usan Round Robin con el quantum "
    "configurado y la cola 2 funciona como FCFS."
)
add_paragraph(doc, "Ventajas:", bold=True)
add_bullets(doc, [
    "Diferencia bien entre tipos de carga.",
    "Permite combinar lo mejor de cada algoritmo.",
])
add_paragraph(doc, "Desventajas:", bold=True)
add_bullets(doc, [
    "La asignación de proceso a cola es fija: si se asigna mal, no hay forma de moverlo después.",
    "Posible inanición de las colas inferiores si las superiores nunca se vacían.",
])

# --- 4.8 MLFQ ---
add_heading(doc, "4.8 Cola multinivel con retroalimentación (MLFQ)", level=2)
add_bullets(doc, [
    ("Tipo", "expropiativo."),
    ("Idea", "igual que la cola multinivel, pero los procesos pueden subir o bajar de cola según su comportamiento."),
])
add_paragraph(doc, "Reglas típicas (y las que usa este simulador):", bold=True)
add_bullets(doc, [
    "Todos los procesos entran en la cola de mayor prioridad.",
    "Si un proceso usa todo su quantum, baja una cola (se le considera CPU-bound).",
    "Si se bloquea por E/S antes de terminar el quantum, mantiene su nivel (se le considera I/O-bound).",
    "Las colas más bajas tienen quantums mayores. En este simulador los valores por defecto son [2, 4, 8].",
])
add_paragraph(doc, "Ventajas:", bold=True)
add_bullets(doc, [
    "Se adapta dinámicamente al comportamiento real del proceso.",
    "Favorece a procesos cortos e interactivos sin necesidad de conocer la ráfaga de antemano.",
    "Es considerado el algoritmo de propósito general más equilibrado.",
])
add_paragraph(doc, "Desventajas:", bold=True)
add_bullets(doc, [
    "Configurar bien el número de colas, los quantums y las reglas de promoción/degradación es difícil.",
    "Mayor complejidad de implementación.",
])
doc.add_page_break()

# =====================================================================
# 5. MEMORIA VIRTUAL
# =====================================================================
add_heading(doc, "5. Conceptos previos de memoria virtual", level=1)
add_bullets(doc, [
    ("Memoria virtual", "abstracción que da a cada proceso un espacio de direcciones propio, generalmente más grande que la RAM disponible."),
    ("Página", "bloque de tamaño fijo del espacio de direcciones virtual."),
    ("Marco (frame)", "bloque de tamaño fijo de la memoria física, del mismo tamaño que una página."),
    ("Tabla de páginas", "estructura que asocia páginas virtuales con marcos físicos."),
    ("Cadena de referencias", "secuencia de páginas que un proceso accede a lo largo del tiempo. Es la entrada típica para evaluar un algoritmo de reemplazo."),
    ("Acierto (hit)", "la página referenciada ya está en algún marco."),
    ("Fallo de página (page fault)", "la página no está en RAM; hay que traerla de disco y, si no hay marcos libres, reemplazar alguna que ya esté."),
    ("Fragmentación interna", "espacio desperdiciado dentro de una página cuando los datos del proceso no llenan completamente el último marco asignado."),
])

add_paragraph(doc,
    "En el simulador la cadena de referencias se genera automáticamente a partir de la "
    "ráfaga y el número de páginas de cada proceso. El módulo de memoria muestra los "
    "marcos en una rejilla y la tabla de páginas por proceso."
)

add_heading(doc, "5.1 Anomalía de Belady", level=2)
add_paragraph(doc,
    "En algunos algoritmos como FIFO puede ocurrir que aumentar el número de marcos "
    "incremente el número de fallos en lugar de reducirlo. Es contraintuitivo y se llama "
    "anomalía de Belady. Los algoritmos que NO la sufren son los llamados algoritmos de "
    "pila (stack algorithms), entre los que están LRU y el óptimo."
)
doc.add_page_break()

# =====================================================================
# 6. REEMPLAZO DE PÁGINAS
# =====================================================================
add_heading(doc, "6. Algoritmos de reemplazo de páginas", level=1)
add_paragraph(doc,
    "El simulador implementa cinco algoritmos en src/engine/memory/. Para los ejemplos "
    "se usa la cadena clásica con tres marcos disponibles: 1, 2, 3, 4, 1, 2, 5, 1, 2, 3, 4, 5."
)

# --- 6.1 FIFO ---
add_heading(doc, "6.1 FIFO (First In First Out)", level=2)
add_bullets(doc, [
    ("Idea", "se reemplaza la página que lleva más tiempo en memoria, sin importar cuándo se usó por última vez."),
    ("Implementación", "cola FIFO de páginas cargadas."),
])
add_paragraph(doc, "Ventajas:", bold=True)
add_bullets(doc, [
    "Sencillo y barato de implementar.",
])
add_paragraph(doc, "Desventajas:", bold=True)
add_bullets(doc, [
    "Puede expulsar páginas que se siguen usando.",
    "Sufre la anomalía de Belady.",
])
add_paragraph(doc, "Sobre la cadena de referencia con 3 marcos genera 9 fallos.", italic=True)

# --- 6.2 LRU ---
add_heading(doc, "6.2 LRU (Least Recently Used)", level=2)
add_bullets(doc, [
    ("Idea", "reemplaza la página que hace más tiempo no se usa. Asume que el pasado reciente predice el futuro cercano (principio de localidad)."),
])
add_paragraph(doc, "Implementaciones típicas:", bold=True)
add_bullets(doc, [
    "Lista enlazada con timestamps: cada acceso mueve la página al frente.",
    "Contador o reloj lógico: en cada acceso se anota el 'tiempo' de uso.",
    "Bits de referencia con muestreo periódico (aproximaciones por hardware).",
])
add_paragraph(doc, "Ventajas:", bold=True)
add_bullets(doc, [
    "Aproxima muy bien al óptimo en cargas con buena localidad.",
    "No sufre la anomalía de Belady (es un algoritmo de pila).",
])
add_paragraph(doc, "Desventajas:", bold=True)
add_bullets(doc, [
    "La implementación pura es costosa: hay que actualizar la estructura en cada acceso.",
    "En hardware, se aproxima con Clock o Segunda oportunidad.",
])

# --- 6.3 Óptimo ---
add_heading(doc, "6.3 Óptimo (OPT o algoritmo de Belady)", level=2)
add_bullets(doc, [
    ("Idea", "reemplaza la página que no se va a usar durante el mayor tiempo en el futuro."),
])
add_paragraph(doc, "Ventajas:", bold=True)
add_bullets(doc, [
    "Es el algoritmo que minimiza el número de fallos. Se usa como cota inferior teórica.",
])
add_paragraph(doc, "Desventajas:", bold=True)
add_bullets(doc, [
    "Irrealizable en la práctica: requiere conocer el futuro.",
    "Solo sirve como referencia para evaluar otros algoritmos.",
])

# --- 6.4 Clock ---
add_heading(doc, "6.4 Reloj (Clock)", level=2)
add_bullets(doc, [
    ("Idea", "aproximación eficiente de LRU."),
])
add_paragraph(doc,
    "Las páginas se organizan en un arreglo circular y cada una tiene un bit de "
    "referencia R. En cada acceso, R se pone a 1. Cuando hace falta reemplazo, un puntero "
    "recorre el círculo: si la página apuntada tiene R = 0 se reemplaza y el puntero "
    "avanza; si tiene R = 1 se pone R = 0 y se avanza al siguiente."
)
add_paragraph(doc, "Ventajas:", bold=True)
add_bullets(doc, [
    "Mucho más barato que LRU puro.",
    "Buena aproximación a LRU en la práctica.",
])
add_paragraph(doc, "Desventajas:", bold=True)
add_bullets(doc, [
    "Si todas las páginas tienen R = 1 puede dar una vuelta completa antes de elegir.",
    "No distingue entre páginas leídas y modificadas. La versión 'Clock mejorado' añade el bit modified.",
])

# --- 6.5 Segunda oportunidad ---
add_heading(doc, "6.5 Segunda oportunidad", level=2)
add_bullets(doc, [
    ("Idea", "variante de FIFO que usa el bit de referencia."),
])
add_paragraph(doc,
    "Cuando hay que reemplazar se mira la página al frente de la cola: si su bit R = 0 "
    "se reemplaza; si R = 1 se le 'perdona': se pone R = 0, se mueve al final de la cola "
    "y se revisa la siguiente."
)
add_paragraph(doc,
    "Conceptualmente, Clock es la versión circular y eficiente de Segunda oportunidad. "
    "Hacen lo mismo, pero Clock evita mover físicamente las páginas dentro de la cola."
)
add_paragraph(doc, "Ventajas:", bold=True)
add_bullets(doc, [
    "Mejora claramente a FIFO sin cambiar mucho su estructura.",
])
add_paragraph(doc, "Desventajas:", bold=True)
add_bullets(doc, [
    "Si todas las páginas tienen R = 1 degenera a FIFO recorriendo toda la cola.",
])
doc.add_page_break()

# =====================================================================
# 7. COMPARATIVAS RÁPIDAS
# =====================================================================
add_heading(doc, "7. Comparativas rápidas", level=1)

add_heading(doc, "7.1 Planificación", level=2)
add_table(doc,
    ["Algoritmo", "Expropiativo", "Necesita ráfaga", "Inanición", "Fortaleza"],
    [
        ["FCFS", "No", "No", "No", "Simplicidad"],
        ["SJF", "No", "Sí", "Sí", "Mínima espera promedio (estática)"],
        ["SRTF", "Sí", "Sí", "Sí", "Mínima espera promedio (dinámica)"],
        ["HRRN", "No", "Sí", "No", "Equilibra cortos y antiguos"],
        ["Round Robin", "Sí", "No", "No", "Tiempo de respuesta"],
        ["Prioridad expropiativa", "Sí", "No", "Sí", "Refleja importancia"],
        ["Cola multinivel", "Depende", "No", "Sí", "Separa clases"],
        ["Cola multinivel con FB", "Sí", "No", "Bajo", "Adaptación dinámica"],
    ])

add_heading(doc, "7.2 Reemplazo de páginas", level=2)
add_table(doc,
    ["Algoritmo", "Costo", "Aprovecha localidad", "Anomalía Belady"],
    [
        ["FIFO", "Muy bajo", "No", "Sí"],
        ["LRU", "Alto (puro)", "Sí", "No"],
        ["Óptimo", "Imposible en la práctica", "Sí (perfecto)", "No"],
        ["Reloj", "Bajo", "Sí (aproximado)", "No"],
        ["Segunda oportunidad", "Bajo", "Sí (aproximado)", "Posible"],
    ])
doc.add_page_break()

# =====================================================================
# 8. ARQUITECTURA TÉCNICA
# =====================================================================
add_heading(doc, "8. Arquitectura técnica del demo", level=1)
add_paragraph(doc,
    "El proyecto separa claramente la lógica de simulación, el estado y la interfaz."
)
add_bullets(doc, [
    ("engine/", "TypeScript puro sin dependencias de React. Contiene los algoritmos y la simulación. Sus index.ts exponen factories como getScheduler y getReplacementAlgorithm para resolver un algoritmo por su nombre."),
    ("store/", "cinco stores de Zustand (procesos, planificación, memoria, tema y tutorial). Actúan como puente entre el motor y la UI. El store del tutorial usa persist para conservar el progreso."),
    ("guide/", "consume tutorialStore y construye un flujo lineal por capítulos y pasos."),
    ("pages/ y components/", "consumen los stores y renderizan la interfaz con React, Tailwind y Framer Motion."),
])

add_heading(doc, "8.1 Flujo de datos", level=2)
add_paragraph(doc,
    "El usuario configura procesos en el formulario. El processStore guarda esa lista. "
    "Al elegir un algoritmo y presionar 'Ejecutar', schedulingStore llama a getScheduler "
    "con el nombre del algoritmo y obtiene una función pura que recibe los procesos y la "
    "configuración. Esa función devuelve la línea de tiempo y las métricas. La UI consume "
    "ese resultado y lo dibuja en el diagrama de Gantt y en las tarjetas de métricas. El "
    "mismo patrón se usa para reemplazo de páginas con getReplacementAlgorithm."
)

add_heading(doc, "8.2 Por qué TypeScript puro en engine/", level=2)
add_bullets(doc, [
    "Permite probar los algoritmos sin levantar la UI.",
    "Hace fácil reusar la lógica en otro contexto (CLI, tests).",
    "Mantiene una frontera clara entre 'cómo se decide' y 'cómo se muestra'.",
])

add_heading(doc, "8.3 Diseño y experiencia", level=2)
add_bullets(doc, [
    ("Mobile-first", "el layout se diseñó primero para celular. El hook useIsMobile decide si se usa Sidebar o MobileTabBar + MobileHeader."),
    ("Tema claro y oscuro", "gestionado por themeStore y aplicado mediante variables CSS."),
    ("Animaciones", "AnimatePresence y motion.div envuelven cada cambio de ruta y cada paso de la guía."),
    ("Componentes reutilizables", "Modal, BottomSheet, Fab, Stepper y StickyActionBar viven en components/ui/."),
])

add_heading(doc, "8.4 Despliegue", level=2)
add_paragraph(doc,
    "El proyecto se despliega automáticamente en Firebase Hosting mediante GitHub Actions. "
    "Al hacer push a main: se instalan dependencias con npm ci, se compila con npm run build "
    "(que corre tsc -b y vite build), y se despliega la carpeta dist/ al proyecto de Firebase. "
    "La configuración firebase.json reescribe todas las rutas a index.html para soportar el "
    "enrutamiento del lado del cliente de React Router."
)
doc.add_page_break()

# =====================================================================
# 9. BANCO DE PREGUNTAS Y RESPUESTAS
# =====================================================================
add_heading(doc, "9. Banco de preguntas y respuestas", level=1)
add_paragraph(doc,
    "Esta sección contiene las preguntas más probables que pueden hacerte. Cada pregunta "
    "tiene una respuesta directa y, cuando aplica, un par de detalles para defenderla."
)

add_heading(doc, "9.1 Preguntas generales del demo", level=2)
add_qa(doc, [
    ("¿Qué es exactamente este proyecto?",
     "Un simulador web de Sistemas Operativos que permite visualizar y experimentar con procesos, planificación de CPU, paginación y reemplazo de páginas. Tiene una guía interactiva por capítulos para alguien sin conocimientos previos y un modo libre con seis módulos para configurar simulaciones completas."),
    ("¿Por qué separaste engine, store y components?",
     "Para que la lógica de los algoritmos no dependa de React. Así puedo probarlos como funciones puras, reutilizarlos y mantener limpio el código de la interfaz. Los stores de Zustand son el puente entre la lógica y la UI."),
    ("¿Por qué Zustand y no Redux?",
     "Porque Zustand es mucho más ligero, no requiere providers ni boilerplate y se integra de forma natural con React 19. Para un proyecto de este tamaño es más que suficiente y permite usar persist de manera trivial."),
    ("¿Cómo se generan las cadenas de referencias en el módulo de memoria?",
     "Automáticamente a partir de la ráfaga y el número de páginas de cada proceso. Cada unidad de ráfaga produce un acceso a alguna de las páginas del proceso, lo que da una secuencia realista para evaluar reemplazo."),
    ("¿La simulación se ejecuta en tiempo real?",
     "Funciona paso a paso o en modo continuo con velocidad ajustable. El motor calcula toda la línea de tiempo de antemano y la UI la 'reproduce' avanzando un tick a la vez."),
    ("¿Soportan multinúcleo?",
     "Sí. Los algoritmos aceptan numCores en la configuración y mantienen un timeline por núcleo. En la UI se ve un Gantt por core."),
])

add_heading(doc, "9.2 Preguntas sobre planificación", level=2)
add_qa(doc, [
    ("¿Qué diferencia hay entre tiempo de espera y tiempo de retorno?",
     "El tiempo de retorno (turnaround) es F - A, todo lo que tarda el proceso desde que llega hasta que termina. El de espera es turnaround menos la ráfaga, o sea solo el tiempo en cola de listos."),
    ("¿Cuándo SJF da el mismo resultado que FCFS?",
     "Cuando los procesos llegan en orden creciente de ráfaga. En ese caso ambos los ejecutan en el mismo orden."),
    ("Demuéstrame que SJF minimiza el tiempo de espera promedio.",
     "Por intercambio: si en una secuencia óptima un proceso largo va antes que uno corto, intercambiarlos no aumenta y normalmente reduce la suma de esperas. Si nunca conviene intercambiar, la secuencia óptima debe estar ordenada por ráfaga creciente, que es lo que hace SJF."),
    ("¿SRTF puede dar peor resultado que SJF?",
     "No en tiempo de espera promedio, pero sí en cambios de contexto, porque SRTF puede expropiar a un proceso a media ráfaga cada vez que llega uno más corto."),
    ("¿Qué pasa si el quantum de Round Robin es 0?",
     "No tiene sentido: la CPU nunca avanzaría y cambiaría de contexto infinitamente. En la práctica el quantum debe ser al menos un tick."),
    ("¿Y si el quantum es enorme?",
     "Round Robin degenera en FCFS porque cada proceso terminaría su ráfaga antes de agotar el quantum."),
    ("¿Cómo se evita la inanición en SJF?",
     "Con aging: aumentar la prioridad efectiva del proceso conforme pasa el tiempo de espera. HRRN ya lo hace de forma natural porque el ratio sube con el tiempo de espera."),
    ("¿Cómo se calcula el response ratio en HRRN?",
     "RR = (espera + ráfaga) / ráfaga. Si la espera es 0 vale 1 (mínimo). Si la espera crece, el RR crece linealmente, así que tarde o temprano cualquier proceso se vuelve el de mayor RR."),
    ("¿Qué es la inversión de prioridad?",
     "Cuando un proceso de alta prioridad queda bloqueado esperando un recurso que tiene uno de baja prioridad, mientras un tercero de prioridad media corre y retrasa al de baja prioridad. Famoso por el caso del Mars Pathfinder. Se mitiga con priority inheritance: el dueño del recurso hereda temporalmente la prioridad del bloqueado."),
    ("¿Qué algoritmo usa Linux?",
     "Históricamente O(1), luego CFS (Completely Fair Scheduler), y desde 2023 EEVDF en versiones recientes. Todos son variantes adaptativas que persiguen el espíritu de MLFQ."),
    ("¿Cuál es la diferencia entre cola multinivel y MLFQ?",
     "En cola multinivel la asignación a una cola es fija. En MLFQ los procesos pueden cambiar de cola según su comportamiento, así que se adapta dinámicamente."),
    ("¿En tu MLFQ qué quantums usas?",
     "Por defecto [2, 4, 8]: la cola más prioritaria tiene quantum 2, la siguiente 4 y la última 8. Quantums crecientes hacia abajo favorecen procesos cortos."),
    ("¿Qué es el efecto convoy?",
     "En FCFS, un proceso CPU-bound largo retiene la CPU y obliga a los procesos cortos a esperar formando un 'convoy'. Es la razón principal por la que FCFS es malo en tiempo de espera."),
    ("¿Cuántos cambios de contexto hace Round Robin con n procesos y ráfaga total T?",
     "Aproximadamente T/q - 1 en el peor caso, donde q es el quantum. Por eso q chico significa más cambios."),
    ("¿FCFS puede dar mejor tiempo de respuesta que SJF?",
     "Sí, si el primer proceso es corto. SJF garantiza el mejor promedio, no necesariamente el mejor en cada caso particular."),
    ("¿Qué algoritmo elegirías para un sistema interactivo?",
     "Round Robin con un quantum pequeño, o MLFQ. La prioridad es minimizar el tiempo de respuesta."),
    ("¿Y para un sistema batch?",
     "FCFS o SJF, donde la prioridad es maximizar throughput y no importa tanto la respuesta inmediata."),
    ("¿Y para un sistema en tiempo real?",
     "Prioridad expropiativa con prioridades estrictas o algoritmos específicos como Rate Monotonic o Earliest Deadline First."),
])

add_heading(doc, "9.3 Preguntas sobre reemplazo de páginas", level=2)
add_qa(doc, [
    ("¿Qué es un fallo de página?",
     "Cuando un proceso accede a una página que no está en RAM. El sistema operativo la carga desde disco y, si no hay marcos libres, reemplaza alguna que ya esté usando el algoritmo de reemplazo."),
    ("¿Qué es la anomalía de Belady?",
     "Que en algunos algoritmos como FIFO, aumentar el número de marcos puede aumentar los fallos en lugar de reducirlos. Es contraintuitivo y por eso preferimos algoritmos de pila como LRU."),
    ("¿Por qué LRU no sufre la anomalía de Belady?",
     "Porque cumple la propiedad de inclusión: el conjunto de páginas residentes con k marcos siempre es subconjunto del conjunto con k+1 marcos. Esa propiedad la tienen los algoritmos llamados 'de pila'."),
    ("¿Cómo se implementa LRU en hardware?",
     "Casi nunca en su forma pura, porque actualizar timestamps en cada acceso es muy caro. Se aproxima con bits de referencia y un puntero giratorio (Clock), o con muestreo periódico de bits."),
    ("¿Cuál es la relación entre Clock y Segunda oportunidad?",
     "Hacen lo mismo conceptualmente: ambos perdonan páginas con bit de referencia 1. Clock es la versión eficiente: en lugar de mover páginas en una cola, las mantiene en un arreglo circular y avanza un puntero."),
    ("¿Por qué el algoritmo Óptimo es solo teórico?",
     "Porque requiere conocer la cadena de referencias completa por adelantado. En tiempo de ejecución no conocemos qué páginas se accederán en el futuro."),
    ("¿Si todos los bits R están en 1 en Clock qué pasa?",
     "El puntero da una vuelta completa, poniendo todos los R a 0. En la siguiente vuelta encuentra una página con R = 0 y la reemplaza. En el peor caso degenera en FIFO con doble pasada."),
    ("¿Cuántos fallos genera FIFO en 1,2,3,4,1,2,5,1,2,3,4,5 con 3 marcos?",
     "9 fallos. Con 4 marcos da 10 (anomalía de Belady)."),
    ("¿Y LRU sobre la misma cadena con 3 marcos?",
     "10 fallos."),
    ("¿Qué es el principio de localidad?",
     "Observación empírica de que los programas tienden a acceder a un conjunto pequeño de páginas durante intervalos de tiempo (localidad temporal y espacial). Es la justificación de LRU: si una página se usó hace poco, probablemente se vuelva a usar."),
    ("¿Qué es el working set?",
     "El conjunto de páginas que un proceso está usando activamente en una ventana de tiempo. Si el sistema mantiene en RAM todos los working sets, hay pocos fallos. Si no, ocurre thrashing."),
    ("¿Qué es thrashing?",
     "Cuando el sistema pasa más tiempo paginando que ejecutando código útil. Pasa cuando hay más procesos compitiendo por memoria que la que cabe sus working sets."),
    ("¿Cómo se mitiga el thrashing?",
     "Reduciendo el grado de multiprogramación: el planificador de mediano plazo suspende procesos hasta que los que quedan tengan suficiente memoria."),
    ("Diferencia entre paginación y segmentación.",
     "La paginación divide la memoria en bloques de tamaño fijo (páginas y marcos) y es transparente al programador. La segmentación divide por unidades lógicas variables (código, pila, datos) y es visible para el programador."),
    ("¿Qué es la fragmentación interna en paginación?",
     "La última página de un proceso casi nunca queda llena. El espacio sobrante dentro de ese marco es fragmentación interna. Es proporcional al tamaño de página."),
    ("Si aumento el tamaño de página, ¿qué pasa?",
     "Disminuye la tabla de páginas y mejora el TLB hit rate, pero aumenta la fragmentación interna y el costo de cada fallo."),
])

add_heading(doc, "9.4 Preguntas técnicas del código", level=2)
add_qa(doc, [
    ("¿Qué hace el archivo metrics.ts?",
     "Define la estructura SchedulableTask, hace deepCopyProcesses para no mutar la entrada, aplana procesos en hilos con flattenTasks y calcula las métricas (tiempos, utilización, cambios de contexto) a partir del timeline producido."),
    ("¿Cómo manejan la expropiación en MLFQ?",
     "Cada tick antes de asignar slots se revisa si una cola de mayor prioridad tiene tareas. Si un slot está corriendo una tarea de cola baja y hay una de cola alta esperando, el slot la libera. La preempted vuelve al frente de su cola sin gastar nivel adicional."),
    ("¿Cómo decide la cola un proceso en multilevelQueue?",
     "Por su prioridad: prioridad 1 va a la cola 0 (más alta), prioridad 2 a la cola 1, y todo lo demás a la cola 2. Esa asignación es fija y nunca cambia."),
    ("¿Cómo calculan utilización de CPU?",
     "Tiempo ocupado dividido entre el tiempo total de la simulación, ambos sumados sobre todos los núcleos."),
    ("¿Y los cambios de contexto?",
     "Cuántas veces el PID en ejecución cambió entre dos slots consecutivos del timeline."),
    ("¿Por qué usaron Vite y no Create React App?",
     "Vite es órdenes de magnitud más rápido en arranque y hot reload, soporta TypeScript de forma nativa y tiene mejor integración con Tailwind 4. Además CRA está deprecated."),
    ("¿Qué hace AppShell?",
     "Es el cascarón visual común al modo libre. En escritorio renderiza Sidebar; en móvil renderiza MobileHeader y MobileTabBar. Le aplica al contenido principal el contenedor con scroll y los espacios reservados para el FAB y la StickyActionBar."),
    ("¿Cómo se persiste el progreso de la guía?",
     "tutorialStore usa el middleware persist de Zustand y guarda el estado en localStorage bajo la clave 'simulador-so-tutorial'. Por eso al recargar el progreso se conserva."),
    ("¿Dónde se ven las animaciones de transición entre páginas?",
     "En App.tsx. Se calcula la dirección (avanzar o retroceder) comparando el índice de la ruta actual contra la anterior y se pasa a AnimatePresence de Framer Motion."),
    ("¿Cómo se garantiza la accesibilidad básica?",
     "Componentes UI con foco visible, aria-labels en botones de la barra inferior, contrastes que respetan el tema, y la barra inferior no se solapa con el contenido gracias al espaciado calculado por la altura de la safe-area."),
])
doc.add_page_break()

# =====================================================================
# 10. ERRORES Y TRAMPAS COMUNES
# =====================================================================
add_heading(doc, "10. Errores y trampas comunes", level=1)
add_numbered(doc, [
    "Confundir tiempo de espera con tiempo de retorno. El de espera es solo tiempo en cola de listos; el de retorno incluye también el tiempo en que el proceso se ejecuta.",
    "Olvidar que SJF y SRTF necesitan conocer la ráfaga de antemano. En sistemas reales se estima con un promedio exponencial.",
    "Asumir que más marcos siempre mejoran FIFO. Falso por la anomalía de Belady.",
    "Tratar Round Robin como justo para tiempo de retorno. Es justo para tiempo de respuesta; el retorno puede ser peor que en SJF.",
    "Confundir prioridad mayor con número mayor. En este simulador, número menor significa prioridad mayor.",
    "No distinguir Clock de Segunda oportunidad. Hacen lo mismo conceptualmente; Clock es la implementación circular y eficiente.",
    "Comparar algoritmos con cargas distintas. Para que la comparación sea válida hay que usar la misma cadena de referencias o el mismo conjunto de procesos.",
    "Pensar que LRU = recencia de creación. LRU mira la última vez que se usó, no cuándo se cargó.",
    "Olvidar contar el primer fallo en cada algoritmo de reemplazo. Cada página cargada por primera vez también cuenta como fallo (compulsory miss).",
    "Dar Round Robin con quantum gigantesco como ejemplo de RR. Con quantum mayor que la mayor ráfaga, es FCFS encubierto.",
])
doc.add_page_break()

# =====================================================================
# 11. CIERRE
# =====================================================================
add_heading(doc, "11. Cierre y conclusiones", level=1)
add_paragraph(doc,
    "El simulador permite tocar y experimentar con conceptos que normalmente solo se "
    "ven en pizarrón. Cada algoritmo tiene su nicho: FCFS por su simplicidad, SJF y SRTF "
    "por minimizar la espera promedio, HRRN por equilibrar cortos y antiguos, Round Robin "
    "por el tiempo de respuesta, prioridad para reflejar importancia, cola multinivel para "
    "separar clases de carga y MLFQ por su capacidad de adaptación dinámica. En reemplazo "
    "de páginas, FIFO es el más simple pero sufre la anomalía de Belady, LRU es el ideal "
    "práctico, OPT es la cota teórica imposible, y Clock y Segunda oportunidad son las "
    "aproximaciones que efectivamente se usan en sistemas reales."
)
add_paragraph(doc,
    "Como mensaje final, no existe un único 'mejor algoritmo'. La elección depende del "
    "tipo de carga y del objetivo (interactividad, throughput, equidad). Por eso el "
    "simulador incluye el módulo de comparación: para que el usuario pueda tomar sus "
    "propias decisiones a partir de los datos."
)
add_paragraph(doc, "Gracias.", bold=True)

# =====================================================================
# GUARDAR
# =====================================================================
out = "Guion_Exposicion_SimuladorSO.docx"
doc.save(out)
print(f"Documento generado: {out}")
